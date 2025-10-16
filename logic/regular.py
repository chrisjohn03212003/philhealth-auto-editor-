# logic/regular.py
from docx import Document
from .utils import find_detailed_table, normalize_and_map_items, money_to_decimal, decimal_to_money
from decimal import Decimal

def process_regular(doc_uploaded: Document, doc_template: Document, filename: str):
    """
    Align uploaded doc to template for 'Regular' category: normalize labels and recompute totals.
    Returns the edited doc (based on template) and output filename.
    """
    out = Document()
    # We'll start from template so formatting matches reference
    out = doc_template

    # Extract detailed table from uploaded doc
    uploaded_table = find_detailed_table(doc_uploaded)
    # Extract rows: assume columns: DATE, REF. NO., QTY, PARTICULARS, UNIT PRICE, DEBIT/CHARGES, ...
    items = []
    if uploaded_table:
        for r in uploaded_table.rows[1:]:
            cells = [c.text.strip() for c in r.cells]
            if len(cells) < 4:
                continue
            particular = normalize_and_map_items(cells[3])
            unit_price = money_to_decimal(cells[4]) if len(cells) > 4 else Decimal('0.00')
            debit = money_to_decimal(cells[5]) if len(cells) > 5 else Decimal('0.00')
            items.append({'particular': particular, 'unit_price': unit_price, 'debit': debit})

    # Aggregate and compute totals
    total = sum(i['debit'] for i in items)
    # Now replace the summary fields in the template doc: search for the text 'TOTAL AMOUNT DUE' or totals near bottom
    # Simple approach: scan paragraphs and replace amounts that match old totals
    # This is heuristic: place primary total into the first numeric placeholder we find after "TOTAL AMOUNT DUE"
    replaced = False
    for p in out.paragraphs:
        if 'TOTAL AMOUNT DUE' in p.text:
            # next paragraphs may contain numbers; try to find and replace the first numeric-looking substring
            # naive: append a new paragraph with correct total
            p_idx = out.paragraphs.index(p)
            out_paragraph = out.paragraphs[p_idx]
            out_paragraph.add_run("\n" + decimal_to_money(total))
            replaced = True
            break

    # (Better: replace exact table cells in the template summary - left as extension)
    output_name = f"edited_regular_{filename}"
    from .utils import ensure_contact_number
    ensure_contact_number(out)
    from .utils import ensure_contact_number
    ensure_contact_number(out)

    return out, output_name
