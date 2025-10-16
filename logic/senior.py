# logic/senior.py
from docx import Document
from .utils import find_detailed_table, normalize_and_map_items, money_to_decimal, decimal_to_money
from decimal import Decimal

def process_senior(doc_uploaded: Document, doc_template: Document, filename: str):
    """
    Process Senior category:
    - apply 20% discount to eligible items (most non-PHIC-covered items)
    - remove VAT if present (this example assumes VAT is not explicitly present in amounts; implement as needed)
    - insert "Senior Citizen Discount (20%)" row into the summary
    """
    out = doc_template
    uploaded_table = find_detailed_table(doc_uploaded)
    items = []
    if uploaded_table:
        for r in uploaded_table.rows[1:]:
            cells = [c.text.strip() for c in r.cells]
            if len(cells) < 6:
                continue
            particular = normalize_and_map_items(cells[3])
            debit = money_to_decimal(cells[5])
            items.append({'particular': particular, 'debit': debit})

    subtotal = sum(i['debit'] for i in items)
    # 20% discount
    discount = (subtotal * Decimal('0.20')).quantize(Decimal('0.01'))
    discounted_total = subtotal - discount

    # find summary area and add the Senior Citizen Discount (20%) row.
    # We'll search for a table containing 'Senior Citizen Discount' text in template; if not found we will append a paragraph.
    found = False
    for table in out.tables:
        for row in table.rows:
            if 'Senior Citizen Discount' in row.cells[0].text:
                # replace second cell with discount shown in parenthesis to match reference format
                row.cells[0].text = 'Senior Citizen Discount (20%)'
                row.cells[1].text = f"({decimal_to_money(discount)})"
                found = True
                break
        if found:
            break

    if not found:
        out.add_paragraph(f"Senior Citizen Discount (20%): ({decimal_to_money(discount)})")

    # Update TOTAL field(s) - heuristic: find 'TOTAL' label in paragraphs and append numeric
    for p in out.paragraphs:
        if 'TOTAL' in p.text.upper():
            p.add_run("\n" + decimal_to_money(discounted_total))
            break

    output_name = f"edited_senior_{filename}"
    return out, output_name
