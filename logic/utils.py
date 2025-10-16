# logic/utils.py
from docx import Document
import re
from decimal import Decimal, ROUND_HALF_UP
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation
import logging
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

def load_docx(path):
    return Document(path)

def save_docx(doc, path):
    doc.save(path)

def decimal_to_money(d):
    """Convert Decimal → string with 2 decimal places and comma separators."""
    if not isinstance(d, Decimal):
        try:
            d = Decimal(str(d))
        except Exception:
            return "0.00"
    return f"{d.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP):,}"

def money_to_decimal(s):
    """Convert money-like strings to Decimal safely."""
    if s is None:
        return Decimal('0.00')

    # Convert to string and clean up
    s = str(s).strip()

    # Empty or just a dash?
    if s == '' or s == '-' or s.lower() in ['n/a', 'none']:
        return Decimal('0.00')

    # Remove commas, peso signs, words, spaces
    s = re.sub(r'[^\d.\-]', '', s)

    # If string is still empty after cleaning → 0.00
    if not s or s == '.' or s == '-':
        return Decimal('0.00')

    try:
        return Decimal(s).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    except InvalidOperation:
        logger.warning(f"[money_to_decimal] Could not parse value: {repr(s)}")
        return Decimal('0.00')

def find_detailed_table(doc):
    # heuristic: find the table that contains "Detailed hospital charges" or has columns like "PARTICULARS", "UNIT PRICE"
    for table in doc.tables:
        first_row_text = ' '.join(cell.text.strip() for cell in table.rows[0].cells)
        if 'PARTICULARS' in first_row_text.upper() or 'UNIT PRICE' in first_row_text.upper():
            return table
    # fallback to first table
    return doc.tables[0] if doc.tables else None

# mapping rules for names we want standardized
STANDARD_NAME_MAP = {
    # map entries found in unedited files to standardized names in references
    # provider id or doctor lines -> Emergent ER Consultation as an example
    r'\d{4}\-\d{7}\-\d-\s*.*': 'Emergent ER Consultation',  # e.g. 1203-2050360-5 DABALOS...
    r'\d{4}\-\d{7}\-\d\s*.*': 'Emergent ER Consultation',
    r'Consultation Fee': 'Emergent ER Consultation',
    # sometimes the detailed line is already Emergent ER Consultation; keep as-is
}

def normalize_and_map_items(particular_text):
    """Return normalized particular label — apply regex-based mapping."""
    txt = particular_text.strip()
    for pattern, replacement in STANDARD_NAME_MAP.items():
        if re.search(pattern, txt, flags=re.IGNORECASE):
            return replacement
    return txt

def ensure_contact_number(doc, contact_text="Contact No.: 8941-8518"):
    """
    Insert 'Contact No.: 8941-8518' right below the paragraph containing 'Accreditation No.'.
    If not found, append it at the end of the document.
    """
    found_contact = False
    found_accreditation = False

    # First, check if contact number already exists
    for p in doc.paragraphs:
        if "contact no" in p.text.lower():
            found_contact = True
            break

    if found_contact:
        # Already has a contact number; do nothing
        return doc

    # Otherwise, look for 'Accreditation No.'
    for i, p in enumerate(doc.paragraphs):
        if "accreditation no" in p.text.lower():
            found_accreditation = True
            # Create a new paragraph element *after* this one
            new_p = OxmlElement("w:p")
            p._element.addnext(new_p)
            # Convert XML paragraph to a docx Paragraph object and add text
            new_para = Paragraph(new_p, p._parent)
            new_para.add_run(contact_text)
            break

    # If we didn't find 'Accreditation No.', just append at the end
    if not found_accreditation:
        doc.add_paragraph(contact_text)

    return doc
